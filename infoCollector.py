from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import matplotlib.pyplot as plt
from matplotlib import font_manager
import os
import argparse
import json

count_word = {
    "编程": 0,
    "语言": 0,
    "程序": 0,
    "算法": 0,
    "数据结构": 0,
    "网络": 0,
    "操作系统": 0,
    "实践": 0,
    "工具": 0,
    "安全": 0,
    "加密": 0,
    "漏洞": 0,
    "思维": 0,
    "模式": 0,
    "数学": 0,
    "硬件": 0,
    "软件": 0,
    "设计": 0,
    "架构": 0,
    "开发": 0,
    "测试": 0,
    "面向对象": 0,
    "函数式": 0,
    "面向过程": 0,
    "并发": 0,
    "分布式": 0,
    "性能": 0,
    "科学": 0,
    "计算": 0,
    "机器学习": 0,
    "深度学习": 0,
    "数据": 0,
    "发布": 0
}

def count_words(content):
    for word in count_word:
        count_word[word] += content.count(word)

def draw_word_statistics():
    plt.figure()
    my_font = font_manager.FontProperties(fname="./SimSun.ttf")
    plt.ylabel("Word")
    plt.xlabel("Count")
    plt.title("Word Statistics")
    plt.yticks(range(len(count_word)), count_word.keys(), fontproperties=my_font, fontsize=7)
    plt.barh(range(len(count_word)), count_word.values())
    plt.savefig("word_statistics.png")

def toDocx(filename: str, folder: str):
    # Get all the json files in the folder
    files = [f for f in os.listdir(folder) if f.endswith(".json")]
    # Create a Document object
    doc = Document()
    for file in files:
        # Open the json file
        with open(os.path.join(folder, file), "r", encoding="utf-8") as f:
            data = json.load(f)
        # Set the title config and write title
        title = doc.add_heading(level=1)
        title_run = title.add_run(data["Title"]) # add title content
        # set title font
        title_run.font.name = u"黑体"
        title_run.element.rPr.rFonts.set(qn("w:eastAsia"), u"黑体")
        title_run.font.size = Pt(24)
        # align title to center
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Set the date config and write date
        updateDate = doc.add_paragraph()
        date_run = updateDate.add_run("更新日期：" + data["Date Update"]) # add date content
        # set date font
        date_run.font.name = u"宋体"
        date_run.element.rPr.rFonts.set(qn("w:eastAsia"), u"宋体")
        date_run.font.size = Pt(16)
        # align date to center
        updateDate.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph() # add a new line
        # Set the description config and write content
        description = doc.add_paragraph()
        description_run = description.add_run(data["Description"]) # add description content
        # count the words in the description
        count_words(data["Description"])
        # set description font
        description_run.font.name = u"宋体"
        description_run.element.rPr.rFonts.set(qn("w:eastAsia"), u"宋体")
        description_run.font.size = Pt(16)
    # Save the docx file
    doc.save(filename)
    # draw the word statistics
    draw_word_statistics()


def toPDF(filename: str, folder: str):
    # Get all the json files in the folder
    files = [f for f in os.listdir(folder) if f.endswith(".json")]
    # Register the font
    pdfmetrics.registerFont(TTFont("simsun", "./SimSun.ttf"))
    # Create a simple template
    pdf = SimpleDocTemplate(filename)
    # Get the default style
    styles = getSampleStyleSheet()
    # Customize the default style
    styles['Normal'].fontName = 'simsun'
    styles['Normal'].fontSize = 12
    styles['Title'].fontName = 'simsun'
    styles['Title'].fontSize = 20
    # Create a list to store the content
    content = []
    for file in files:
        # Open the json file
        with open(os.path.join(folder, file), "r", encoding="utf-8") as f:
            data = json.load(f)
        # Add title
        content.append(Paragraph(data["Title"], styles["Title"]))
        # Add date
        content.append(Paragraph("更新日期：" + data["Date Update"], styles["Normal"]))
        # Add description
        content.append(Paragraph(data["Description"], styles["Normal"]))
        # count the words in the description
        count_words(data["Description"])
    # Build the pdf file
    pdf.build(content)
    # draw the word statistics
    draw_word_statistics()

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Convert JSON to DOCX")
    parser.add_argument("-f", "--folder", help="Folder path", required=True)
    parser.add_argument("-o", "--output", help="Output file", required=True)
    args = parser.parse_args()

    ext = os.path.splitext(args.output)[1]
    if ext == ".docx":
        toDocx(args.output, args.folder)
    elif ext == ".pdf":
        toPDF(args.output, args.folder)
    else:
        print("Unsupported file format")